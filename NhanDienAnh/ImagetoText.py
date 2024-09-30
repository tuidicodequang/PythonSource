import os
import cv2
import numpy as np
from google.cloud import vision
from google.cloud.vision_v1 import types
import io
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

# Đảm bảo bạn đã thiết lập biến môi trường GOOGLE_APPLICATION_CREDENTIALS

def advanced_preprocess(image_path):
    img = cv2.imread(image_path)
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    
    # Khử nhiễu bằng Non-local Means Denoising
    denoised = cv2.fastNlMeansDenoising(gray, None, 10, 7, 21)
    
    # Tăng cường độ tương phản bằng CLAHE
    clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8))
    contrast = clahe.apply(denoised)
    
    # Áp dụng Adaptive Thresholding
    binary = cv2.adaptiveThreshold(contrast, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2)
    
    # Áp dụng morphological operations
    kernel = np.ones((1,1), np.uint8)
    opening = cv2.morphologyEx(binary, cv2.MORPH_OPEN, kernel, iterations=1)
    closing = cv2.morphologyEx(opening, cv2.MORPH_CLOSE, kernel, iterations=1)
    
    # Khử nhiễu một lần nữa sau các phép biến đổi
    final = cv2.fastNlMeansDenoising(closing, None, 10, 7, 21)
    
    return final

def detect_text_google_vision(image_path):
    client = vision.ImageAnnotatorClient()

    with io.open(image_path, 'rb') as image_file:
        content = image_file.read()

    image = types.Image(content=content)
    response = client.text_detection(image=image)
    texts = response.text_annotations

    if texts:
        return texts[0].description
    return ""

def clean_text(text):
    # Loại bỏ các ký tự không mong muốn nhưng giữ nguyên khoảng trắng và ngắt dòng
    text = re.sub(r'[^\w\s\.,?!-]', '', text)
    # Chuẩn hóa khoảng trắng (không xóa ngắt dòng)
    text = re.sub(r' +', ' ', text)
    # Loại bỏ khoảng trắng ở đầu và cuối mỗi dòng
    text = '\n'.join(line.strip() for line in text.splitlines())
    return text

def add_text_to_word(doc, text, filename):
    doc.add_heading(f"Image: {filename}", level=1)
    
    for paragraph in text.split('\n'):
        p = doc.add_paragraph(paragraph)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
    
    doc.add_page_break()

def process_images_in_folder(image_folder, output_word):
    doc = Document()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    for filename in os.listdir(image_folder):
        if filename.lower().endswith(('.png', '.jpg', '.jpeg')):
            image_path = os.path.join(image_folder, filename)
            print(f"Đang xử lý ảnh: {filename}")
            
            # Tiền xử lý ảnh
            processed_image = advanced_preprocess(image_path)
            
            # Lưu ảnh đã xử lý
            processed_path = os.path.join(image_folder, f"processed_{filename}")
            cv2.imwrite(processed_path, processed_image)
            
            # Sử dụng Google Vision API để nhận dạng văn bản
            text = detect_text_google_vision(processed_path)
            
            # Làm sạch và định dạng văn bản
            text = clean_text(text)
            
            # Thêm vào tài liệu Word
            add_text_to_word(doc, text, filename)
            
            # Xóa ảnh đã xử lý
            os.remove(processed_path)

    doc.save(output_word)
    print(f"File Word đã được lưu tại: {output_word}")

# Đường dẫn thư mục chứa ảnh
image_folder = r'C:\Users\manno\OneDrive\Desktop\New folder'
output_word = 'output_document_advanced.docx'

process_images_in_folder(image_folder, output_word)